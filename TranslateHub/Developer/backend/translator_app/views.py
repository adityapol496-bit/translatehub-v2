import io
import json

from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.http import HttpResponse, JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

from .models import ChatSession, Translation, UserProfile

try:
    from deep_translator import GoogleTranslator
except ImportError:  # pragma: no cover
    GoogleTranslator = None

try:
    from langdetect import detect as _detect_lang
except ImportError:  # pragma: no cover
    _detect_lang = None

try:
    from gtts import gTTS
except ImportError:  # pragma: no cover
    gTTS = None

try:
    import pytesseract
    from PIL import Image
except ImportError:  # pragma: no cover
    pytesseract = None
    Image = None

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

try:
    from indic_transliteration import sanscript
    from indic_transliteration.sanscript import transliterate as _indic_transliterate
except ImportError:  # pragma: no cover
    sanscript = None
    _indic_transliterate = None

# Maps our target_lang codes to the indic_transliteration script name for
# languages that use a non-Latin Indic script. Languages not listed here
# either already use Latin script or aren't supported by the library, so no
# transliteration line is generated for them.
_INDIC_SCRIPT_MAP = {
    "mr": "devanagari",
    "hi": "devanagari",
    "ne": "devanagari",
    "sa": "devanagari",
    "kok": "devanagari",
    "mai": "devanagari",
    "doi": "devanagari",
    "bho": "devanagari",
    "kn": "kannada",
    "ta": "tamil",
    "te": "telugu",
    "gu": "gujarati",
    "pa": "gurmukhi",
    "bn": "bengali",
    "as": "bengali",
    "ml": "malayalam",
    "or": "oriya",
}


def _get_transliteration(text, target_lang):
    """Best-effort romanization of Indic-script translations, e.g. so the UI
    can show 'reel-time bhashantar' under 'रिअल-टाइम भाषांतर'. Returns None
    for scripts we don't have a mapping for (Latin, CJK, Cyrillic, etc.)."""
    script_name = _INDIC_SCRIPT_MAP.get(target_lang)
    if not script_name or _indic_transliterate is None or not text:
        return None
    try:
        result = _indic_transliterate(text, getattr(sanscript, script_name.upper()), sanscript.ITRANS)
        return result.lower()
    except Exception:  # noqa: BLE001 — transliteration is a nice-to-have, never fail the request over it
        return None


def _get_or_create_session(session_id):
    if not session_id:
        return None
    session, _ = ChatSession.objects.get_or_create(session_id=session_id)
    return session


@csrf_exempt
@require_http_methods(["POST"])
def ocr_translate_view(request):
    """POST /api/ocr-translate/  multipart form:
        image: <file>
        target_lang: "mr" (or any language code)
    Extracts text from an uploaded photo (poster/sign/document) using OCR,
    then translates it to the target language.
    """
    if pytesseract is None or Image is None:
        return JsonResponse(
            {"error": "OCR is not installed on the server (pytesseract/Pillow missing)."},
            status=500,
        )

    image_file = request.FILES.get("image")
    if not image_file:
        return JsonResponse({"error": "No image uploaded."}, status=400)

    target_lang = request.POST.get("target_lang") or "en"
    session_id = request.POST.get("session_id")

    try:
        img = Image.open(image_file)
        extracted_text = pytesseract.image_to_string(img).strip()
    except Exception as exc:
        return JsonResponse({"error": f"Could not read text from image: {exc}"}, status=500)

    if not extracted_text:
        return JsonResponse({"error": "No readable text found in that image."}, status=422)

    if GoogleTranslator is None:
        return JsonResponse({"error": "Translator not available on server."}, status=500)

    try:
        translated_text = GoogleTranslator(source="auto", target=target_lang).translate(extracted_text)
    except Exception as exc:
        return JsonResponse({"error": f"Translation failed: {exc}"}, status=500)

    session = _get_or_create_session(session_id)
    Translation.objects.create(
        user=request.user if request.user.is_authenticated else None,
        session=session,
        source_text=extracted_text,
        translated_text=translated_text or "",
        source_lang="auto",
        target_lang=target_lang,
    )

    return JsonResponse({
        "extracted_text": extracted_text,
        "translated_text": translated_text,
    })


@csrf_exempt
@require_http_methods(["POST"])
def translate_view(request):
    """
    POST /api/translate/
    body: { "text": "...", "source_lang": "en" | "detect", "target_lang": "mr", "session_id": "optional" }
    returns: { "translated_text": "...", "detected_lang": "en" }
    If the user is logged in (session cookie), the translation is also linked to their account.
    """
    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON body."}, status=400)

    text = (data.get("text") or "").strip()
    source_lang = data.get("source_lang") or "auto"
    target_lang = data.get("target_lang") or "en"
    session_id = data.get("session_id")

    if not text:
        return JsonResponse({"error": "Field 'text' is required."}, status=400)

    if GoogleTranslator is None:
        return JsonResponse(
            {"error": "deep-translator is not installed. Run: pip install deep-translator"},
            status=500,
        )

    # The frontend sends "detect" for auto-detection; deep-translator expects "auto".
    dt_source = "auto" if source_lang in ("detect", "auto", "") else source_lang

    try:
        translator = GoogleTranslator(source=dt_source, target=target_lang)
        try:
            translated_text = translator.translate(text)
        except Exception:
            # deep-translator hits the live Google Translate site, which
            # occasionally has a transient hiccup — one retry clears most of those.
            translated_text = translator.translate(text)
    except Exception as exc:  # noqa: BLE001 — surface translator errors to the client
        return JsonResponse({"error": f"Translation failed: {exc}"}, status=502)

    session = _get_or_create_session(session_id)
    Translation.objects.create(
        user=request.user if request.user.is_authenticated else None,
        session=session,
        source_text=text,
        translated_text=translated_text or "",
        source_lang=source_lang,
        target_lang=target_lang,
    )

    transliteration = _get_transliteration(translated_text, target_lang)

    detected_lang = source_lang
    if dt_source == "auto" and _detect_lang is not None:
        try:
            detected_lang = _detect_lang(text)
        except Exception:
            detected_lang = source_lang

    return JsonResponse({
        "translated_text": translated_text,
        "transliteration": transliteration,
        "detected_lang": detected_lang,
    })


@require_http_methods(["GET"])
def history_view(request):
    """
    GET /api/history/?session_id=...
    Returns translations for the logged-in user (if authenticated), otherwise
    falls back to the anonymous browser session_id.
    """
    session_id = request.GET.get("session_id")

    if request.user.is_authenticated:
        qs = Translation.objects.filter(user=request.user)
    elif session_id:
        qs = Translation.objects.filter(session__session_id=session_id)
    else:
        qs = Translation.objects.none()

    qs = qs[:50]
    results = [
        {
            "id": t.id,
            "source_text": t.source_text,
            "translated_text": t.translated_text,
            "source_lang": t.source_lang,
            "target_lang": t.target_lang,
            "created_at": t.created_at.isoformat(),
        }
        for t in qs
    ]
    return JsonResponse({"results": results})


@csrf_exempt
@require_http_methods(["POST"])
def clear_history_view(request):
    """POST /api/history/clear/  body: { "session_id": "optional" } """
    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        data = {}

    session_id = data.get("session_id")

    if request.user.is_authenticated:
        qs = Translation.objects.filter(user=request.user)
    elif session_id:
        qs = Translation.objects.filter(session__session_id=session_id)
    else:
        qs = Translation.objects.none()

    deleted_count, _ = qs.delete()
    return JsonResponse({"deleted": deleted_count})


@csrf_exempt
@require_http_methods(["POST", "DELETE"])
def delete_history_item_view(request, translation_id):
    """POST or DELETE /api/history/<id>/delete/  body: { "session_id": "optional" }
    Deletes a single translation — only if it belongs to the logged-in user,
    or (for anonymous visitors) to the browser session_id that owns it.
    """
    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        data = {}

    session_id = data.get("session_id") or request.GET.get("session_id")

    if request.user.is_authenticated:
        qs = Translation.objects.filter(id=translation_id, user=request.user)
    elif session_id:
        qs = Translation.objects.filter(id=translation_id, session__session_id=session_id)
    else:
        qs = Translation.objects.none()

    deleted_count, _ = qs.delete()
    if not deleted_count:
        return JsonResponse({"error": "Translation not found."}, status=404)
    return JsonResponse({"deleted": deleted_count})


@require_http_methods(["GET"])
def download_history_item_view(request, translation_id):
    """GET /api/history/<id>/download/?session_id=optional&format=docx
    Downloads a single translation entry as a Word (.docx) file.
    Only accessible to the owner (logged-in user, or matching anonymous session_id).
    """
    session_id = request.GET.get("session_id")

    if request.user.is_authenticated:
        qs = Translation.objects.filter(id=translation_id, user=request.user)
    elif session_id:
        qs = Translation.objects.filter(id=translation_id, session__session_id=session_id)
    else:
        qs = Translation.objects.none()

    t = qs.first()
    if not t:
        return JsonResponse({"error": "Translation not found."}, status=404)

    if Document is None:
        return JsonResponse({"error": "python-docx is not installed on the server."}, status=500)

    doc = Document()
    doc.add_heading("TranslateHub — Translation", level=1)
    doc.add_paragraph(f"Date: {t.created_at.strftime('%d %b %Y, %I:%M %p')}")
    doc.add_paragraph(f"From: {t.source_lang}   →   To: {t.target_lang}")
    doc.add_paragraph("")
    doc.add_heading("Original Text", level=2)
    doc.add_paragraph(t.source_text)
    doc.add_heading("Translated Text", level=2)
    doc.add_paragraph(t.translated_text)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="translation_{t.id}.docx"'
    return response


# ---------------------------------------------------------------------------
# AUTH: signup, login, logout, profile
# ---------------------------------------------------------------------------

def _user_payload(user):
    profile = getattr(user, "profile", None)
    return {
        "username": user.username,
        "name": user.first_name or user.username,
        "email": user.email,
        "phone": profile.phone if profile else "",
        "preferred_lang": profile.preferred_lang if profile else "en",
        "photo": profile.photo_base64 if profile else "",
    }


@csrf_exempt
@require_http_methods(["POST"])
def forgot_password_view(request):
    """POST /api/auth/forgot-password/  body: { "email" }
    Generates a 6-digit OTP, saves it against the user's profile (valid for
    10 minutes), and emails it to them."""
    import random
    from datetime import timedelta
    from django.utils import timezone
    from django.core.mail import send_mail

    body = json.loads(request.body or "{}")
    email = (body.get("email") or "").strip().lower()

    user = User.objects.filter(username=email).first()
    if user:
        profile, _ = UserProfile.objects.get_or_create(user=user)
        otp = f"{random.randint(0, 999999):06d}"
        profile.otp_code = otp
        profile.otp_expires_at = timezone.now() + timedelta(minutes=10)
        profile.save()

        try:
            send_mail(
                subject="TranslateHub — Your OTP code",
                message=(
                    f"Your TranslateHub verification code is: {otp}\n\n"
                    "This code expires in 10 minutes. If you didn't request "
                    "this, you can safely ignore this email."
                ),
                from_email=None,  # uses DEFAULT_FROM_EMAIL
                recipient_list=[email],
                fail_silently=False,
            )
        except Exception as exc:
            # Don't leak email-server errors to the client, but do log them
            # server-side (visible in Railway's Deploy Logs) so a real
            # delivery problem (bad SMTP credentials, etc.) is easy to spot.
            print(f"[forgot_password_view] Failed to send OTP email to {email}: {exc}")

    # Always respond the same way whether or not the email exists, so this
    # endpoint can't be used to check which emails are registered.
    return JsonResponse(
        {"message": "If an account exists for that email, an OTP has been sent."}
    )


@csrf_exempt
@require_http_methods(["POST"])
def verify_otp_view(request):
    """POST /api/auth/verify-otp/  body: { "email", "otp" } """
    from django.utils import timezone

    body = json.loads(request.body or "{}")
    email = (body.get("email") or "").strip().lower()
    otp = (body.get("otp") or "").strip()

    user = User.objects.filter(username=email).first()
    profile = getattr(user, "profile", None) if user else None

    if not profile or not profile.otp_code or profile.otp_code != otp:
        return JsonResponse({"error": "Invalid OTP."}, status=400)
    if not profile.otp_expires_at or profile.otp_expires_at < timezone.now():
        return JsonResponse({"error": "OTP has expired. Please request a new one."}, status=400)

    return JsonResponse({"message": "OTP verified."})


@csrf_exempt
@require_http_methods(["POST"])
def reset_password_view(request):
    """POST /api/auth/reset-password/  body: { "email", "otp", "new_password" }
    Re-validates the OTP, then sets the new password and clears the OTP."""
    from django.utils import timezone

    body = json.loads(request.body or "{}")
    email = (body.get("email") or "").strip().lower()
    otp = (body.get("otp") or "").strip()
    new_password = body.get("new_password") or ""

    if len(new_password) < 6:
        return JsonResponse({"error": "Password must be at least 6 characters."}, status=400)

    user = User.objects.filter(username=email).first()
    profile = getattr(user, "profile", None) if user else None

    if not profile or not profile.otp_code or profile.otp_code != otp:
        return JsonResponse({"error": "Invalid OTP."}, status=400)
    if not profile.otp_expires_at or profile.otp_expires_at < timezone.now():
        return JsonResponse({"error": "OTP has expired. Please request a new one."}, status=400)

    user.set_password(new_password)
    user.save()
    profile.otp_code = ""
    profile.otp_expires_at = None
    profile.save()

    return JsonResponse({"message": "Password reset successfully."})


@csrf_exempt
@require_http_methods(["POST"])
def signup_view(request):
    """POST /api/auth/signup/  body: { "name", "email", "password" } """
    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON body."}, status=400)

    name = (data.get("name") or "").strip()
    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""

    if not email or not password:
        return JsonResponse({"error": "Email and password are required."}, status=400)

    if User.objects.filter(username=email).exists():
        return JsonResponse({"error": "An account with this email already exists."}, status=409)

    user = User.objects.create_user(username=email, email=email, password=password, first_name=name)
    UserProfile.objects.create(user=user)

    login(request, user)
    return JsonResponse({"user": _user_payload(user)})


@csrf_exempt
@require_http_methods(["POST"])
def login_view(request):
    """POST /api/auth/login/  body: { "email", "password" } """
    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON body."}, status=400)

    email = (data.get("email") or "").strip().lower()
    password = data.get("password") or ""

    user = authenticate(request, username=email, password=password)
    if user is None:
        return JsonResponse({"error": "Invalid email or password."}, status=401)

    login(request, user)
    return JsonResponse({"user": _user_payload(user)})


@csrf_exempt
@require_http_methods(["POST"])
def google_login_view(request):
    """POST /api/auth/google/  body: { "credential": "<Google ID token>" } """
    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON body."}, status=400)

    token = data.get("credential")
    if not token:
        return JsonResponse({"error": "Missing Google credential."}, status=400)

    try:
        from google.oauth2 import id_token as google_id_token
        from google.auth.transport import requests as google_requests
        from django.conf import settings

        idinfo = google_id_token.verify_oauth2_token(
            token, google_requests.Request(), settings.GOOGLE_CLIENT_ID
        )
    except Exception:
        return JsonResponse({"error": "Invalid Google token."}, status=401)

    email = (idinfo.get("email") or "").strip().lower()
    name = idinfo.get("name") or email.split("@")[0]

    if not email:
        return JsonResponse({"error": "Google account has no email."}, status=400)

    user, created = User.objects.get_or_create(
        username=email, defaults={"email": email, "first_name": name}
    )
    if created:
        user.set_unusable_password()
        user.save()
        UserProfile.objects.create(user=user)

    login(request, user)
    return JsonResponse({"user": _user_payload(user)})


@require_http_methods(["POST"])
def logout_view(request):
    """POST /api/auth/logout/"""
    logout(request)
    return JsonResponse({"success": True})


@require_http_methods(["GET"])
def me_view(request):
    """GET /api/auth/me/ — returns the logged-in user, or null if not authenticated."""
    if not request.user.is_authenticated:
        return JsonResponse({"user": None})
    return JsonResponse({"user": _user_payload(request.user)})


@csrf_exempt
@require_http_methods(["GET", "POST"])
def profile_view(request):
    """
    GET  /api/auth/profile/  → current profile
    POST /api/auth/profile/  body: { "name", "phone", "preferred_lang" }
    Requires the user to be logged in.
    """
    if not request.user.is_authenticated:
        return JsonResponse({"error": "You must be logged in."}, status=401)

    profile, _ = UserProfile.objects.get_or_create(user=request.user)

    if request.method == "POST":
        try:
            data = json.loads(request.body or "{}")
        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON body."}, status=400)

        request.user.first_name = data.get("name", request.user.first_name)
        request.user.save()
        profile.phone = data.get("phone", profile.phone)
        profile.preferred_lang = data.get("preferred_lang", profile.preferred_lang)
        if "photo" in data:
            profile.photo_base64 = data.get("photo") or ""
        profile.save()

    return JsonResponse({"user": _user_payload(request.user)})


@csrf_exempt
@require_http_methods(["POST"])
def speak_view(request):
    """POST /api/speak/  body: { "text": "...", "lang": "mr" }
    Returns an MP3 audio stream of the text spoken in the given language
    using Google's TTS (via gTTS) — works for every language, no matter
    what voices are installed on the user's computer/browser.
    """
    if gTTS is None:
        return JsonResponse({"error": "gTTS is not installed on the server."}, status=500)

    try:
        data = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON body."}, status=400)

    text = (data.get("text") or "").strip()
    lang = (data.get("lang") or "en").strip()

    if not text:
        return JsonResponse({"error": "No text provided."}, status=400)

    try:
        buffer = io.BytesIO()
        gTTS(text=text, lang=lang).write_to_fp(buffer)
        buffer.seek(0)
        return HttpResponse(buffer.read(), content_type="audio/mpeg")
    except Exception as exc:
        return JsonResponse({"error": f"Could not generate speech: {exc}"}, status=500)
